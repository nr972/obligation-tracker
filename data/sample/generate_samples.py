"""Generate synthetic sample contract DOCX files for demo/testing.

Run this script once to create the sample contracts:
    python data/sample/generate_samples.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

CONTRACTS_DIR = Path(__file__).parent / "contracts"
CONTRACTS_DIR.mkdir(exist_ok=True)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def create_saas_agreement():
    doc = Document()
    doc.add_heading("CloudSync Platform Service Agreement", 0)
    add_para(doc, "This Service Agreement (\"Agreement\") is entered into as of January 15, 2025, by and between Nimbus Technologies Inc. (\"Provider\") and Client (\"Client\").")

    add_heading(doc, "1. Definitions")
    add_para(doc, "\"Platform\" means the CloudSync cloud-based data synchronization and management platform.")
    add_para(doc, "\"Service Credits\" means credits issued to Client as a result of Provider's failure to meet the SLA commitments.")

    add_heading(doc, "2. Term")
    add_para(doc, "This Agreement is effective from January 15, 2025 through January 14, 2027. The Agreement will automatically renew for successive one-year periods unless either party provides 90 days written notice of non-renewal.")

    add_heading(doc, "3. Service Level Agreement")
    add_para(doc, "3.1 Provider guarantees 99.9% platform availability measured on a monthly basis, excluding scheduled maintenance windows communicated at least 72 hours in advance.")
    add_para(doc, "3.2 For each 0.1% below 99.9% uptime, Client receives a 10% service credit on the monthly fee, up to a maximum of 30% of monthly fees.")
    add_para(doc, "3.3 Provider shall deliver monthly SLA performance reports within 10 business days of each calendar month end, detailing uptime metrics, incident response times, and resolution statistics. Failure to deliver reports for two consecutive months entitles Client to a 5% service credit.")

    add_heading(doc, "4. Payment Terms")
    add_para(doc, "4.1 Client shall pay the quarterly subscription fee of $45,000 within 30 days of invoice receipt.")
    add_para(doc, "4.2 Late payment incurs 1.5% monthly interest on the outstanding balance.")

    add_heading(doc, "5. Data Protection")
    add_para(doc, "5.1 Provider shall perform daily backups of all Client data and maintain backup copies for a minimum of 30 days. Data loss due to backup failure subjects Provider to liability for direct damages.")
    add_para(doc, "5.2 Provider shall notify Client of any security breach affecting Client data within 72 hours of discovery. Failure to provide timely notification constitutes a material breach of this Agreement.")

    add_heading(doc, "6. Security")
    add_para(doc, "6.1 Provider shall conduct an independent third-party security audit annually and share the SOC 2 Type II report with Client within 30 days of completion.")
    add_para(doc, "6.2 Provider shall maintain industry-standard encryption for data in transit and at rest.")

    add_heading(doc, "7. Confidentiality")
    add_para(doc, "Both parties agree to maintain the confidentiality of all proprietary information disclosed under this Agreement for the term plus three years.")

    add_heading(doc, "8. Termination")
    add_para(doc, "Either party may terminate this Agreement for material breach upon 30 days written notice if the breach is not cured within the notice period.")

    add_heading(doc, "9. General Provisions")
    add_para(doc, "This Agreement shall be governed by the laws of the State of Delaware.")

    doc.save(CONTRACTS_DIR / "saas_agreement_nimbus.docx")


def create_vendor_agreement():
    doc = Document()
    doc.add_heading("IT Infrastructure Support Agreement", 0)
    add_para(doc, "This Agreement is entered into as of March 1, 2025, by and between Apex Solutions Group (\"Vendor\") and Client (\"Client\").")

    add_heading(doc, "1. Scope of Services")
    add_para(doc, "Vendor shall provide IT infrastructure support services including network management, server maintenance, and help desk support.")

    add_heading(doc, "2. Term")
    add_para(doc, "This Agreement is effective from March 1, 2025 through February 28, 2026. Renewal requires mutual written agreement at least 60 days before expiration.")

    add_heading(doc, "3. Governance")
    add_para(doc, "Both parties shall participate in quarterly business reviews to assess service delivery, discuss upcoming needs, and review performance metrics.")

    add_heading(doc, "4. Insurance")
    add_para(doc, "Vendor shall maintain professional liability insurance of no less than $2,000,000 and provide proof of coverage upon request. Failure to maintain required insurance is grounds for immediate termination.")

    add_heading(doc, "5. Personnel")
    add_para(doc, "Vendor shall conduct background checks on all personnel assigned to Client facilities or with access to Client systems prior to assignment.")

    add_heading(doc, "6. Payment Terms")
    add_para(doc, "Vendor shall submit invoices within 15 days of service delivery. Client shall pay within 45 days of invoice receipt. Late payment incurs 1% monthly interest.")

    add_heading(doc, "7. Termination")
    add_para(doc, "Upon termination, Vendor shall return all Client-provided equipment within 10 business days. Client may deduct equipment replacement costs from final payment.")

    add_heading(doc, "8. Confidentiality")
    add_para(doc, "Vendor shall maintain the confidentiality of all Client information encountered during service delivery.")

    doc.save(CONTRACTS_DIR / "vendor_agreement_apex.docx")


def create_dpa():
    doc = Document()
    doc.add_heading("Data Processing Addendum", 0)
    add_para(doc, "This Data Processing Addendum (\"DPA\") supplements the Master Services Agreement between DataVault Analytics LLC (\"Processor\") and Client (\"Controller\"), effective February 1, 2025.")

    add_heading(doc, "1. Definitions")
    add_para(doc, "Terms used herein have the meanings set forth in applicable data protection legislation, including GDPR and CCPA as applicable.")

    add_heading(doc, "2. Processing Scope")
    add_para(doc, "Processor shall process personal data solely as instructed by Controller and for the purposes described in Annex A.")

    add_heading(doc, "3. Security Measures")
    add_para(doc, "Processor shall implement and maintain appropriate technical and organizational measures to ensure a level of security appropriate to the risk of processing, including encryption, access controls, and regular security testing.")

    add_heading(doc, "4. Sub-Processors")
    add_para(doc, "Processor shall provide Controller with 30 days prior written notice before engaging any new sub-processor, including the sub-processor's identity and the processing activities. Engaging a sub-processor without proper notice constitutes a material breach.")

    add_heading(doc, "5. Data Protection Impact Assessment")
    add_para(doc, "Processor shall assist Controller in conducting Data Protection Impact Assessments upon written request, providing all necessary information within 15 business days.")

    add_heading(doc, "6. Data Breach")
    add_para(doc, "Processor shall notify Controller of any personal data breach within 48 hours of becoming aware, providing details required under applicable data protection law. Delayed notification may result in regulatory fines for which Processor shall indemnify Controller.")

    add_heading(doc, "7. Compliance")
    add_para(doc, "Processor shall provide annual written certification of compliance with all obligations under this DPA, including technical and organizational measures.")

    add_heading(doc, "8. Termination")
    add_para(doc, "Upon termination, Processor shall delete or return all personal data within 30 days and provide written certification of deletion. Failure to delete data subjects Processor to regulatory penalties and indemnification obligations.")

    add_heading(doc, "9. Term")
    add_para(doc, "This DPA remains in effect from February 1, 2025 through January 31, 2027, and automatically renews unless either party provides 30 days written notice.")

    doc.save(CONTRACTS_DIR / "dpa_datavault.docx")


def create_consulting_agreement():
    doc = Document()
    doc.add_heading("Strategic Advisory Services Agreement", 0)
    add_para(doc, "This Agreement is entered into as of April 1, 2025, by and between Meridian Consulting Partners (\"Consultant\") and Client (\"Client\").")

    add_heading(doc, "1. Scope of Engagement")
    add_para(doc, "Consultant shall provide strategic advisory services related to legal technology transformation and operational efficiency.")

    add_heading(doc, "2. Term")
    add_para(doc, "This Agreement is effective from April 1, 2025 through December 31, 2025. No automatic renewal.")

    add_heading(doc, "3. Deliverables")
    add_para(doc, "Consultant shall submit a written report of deliverables and progress by the 5th business day of each month.")

    add_heading(doc, "4. Intellectual Property")
    add_para(doc, "All work product created by Consultant in the course of this engagement shall be assigned to Client upon creation as work-made-for-hire. Failure to assign IP constitutes a material breach.")

    add_heading(doc, "5. Restrictive Covenants")
    add_para(doc, "Consultant shall not provide substantially similar advisory services to Client's direct competitors for 12 months following termination of this Agreement.")

    add_heading(doc, "6. Payment Terms")
    add_para(doc, "Consultant shall submit invoices bi-weekly. Client shall pay within 15 business days. Late payment incurs 1.5% monthly interest.")

    add_heading(doc, "7. Expenses")
    add_para(doc, "Any expense exceeding $500 must receive prior written approval from Client's project lead before being incurred. Approved expenses will be reimbursed within 30 days of submission with receipts.")

    add_heading(doc, "8. Confidentiality")
    add_para(doc, "Consultant shall maintain strict confidentiality of all Client information for the term plus two years.")

    doc.save(CONTRACTS_DIR / "consulting_meridian.docx")


def create_license_agreement():
    doc = Document()
    doc.add_heading("Enterprise Software License Agreement", 0)
    add_para(doc, "This License Agreement is entered into as of January 1, 2025, by and between Prism Software Corp. (\"Licensor\") and Client (\"Licensee\").")

    add_heading(doc, "1. License Grant")
    add_para(doc, "Licensor grants Licensee a non-exclusive, non-transferable license to use the software for internal business purposes for up to 500 named users.")

    add_heading(doc, "2. Term")
    add_para(doc, "This Agreement is effective from January 1, 2025 through December 31, 2026. The Agreement automatically renews for successive one-year periods unless either party provides 60 days written notice of non-renewal.")

    add_heading(doc, "3. License Fees")
    add_para(doc, "Licensee shall pay the annual license fee of $120,000 within 30 days of each anniversary date. Non-payment within 60 days entitles Licensor to suspend access to the software.")

    add_heading(doc, "4. Audit Rights")
    add_para(doc, "Licensor reserves the right to audit Licensee's usage annually. Licensee shall provide access to usage records within 10 business days of audit request. Non-compliance may result in license termination.")

    add_heading(doc, "5. Compliance")
    add_para(doc, "Licensee shall comply with all applicable export control laws and regulations in its use of the software. Violation may result in immediate termination and legal action.")

    add_heading(doc, "6. Security")
    add_para(doc, "Licensee shall maintain access logs for all users of the licensed software for a minimum of 12 months.")

    add_heading(doc, "7. Maintenance")
    add_para(doc, "Licensee shall update to the latest supported version of the software within 90 days of release notification. Running unsupported versions voids the support warranty.")

    add_heading(doc, "8. Limitation of Liability")
    add_para(doc, "In no event shall either party's total liability exceed the fees paid in the 12 months preceding the claim.")

    doc.save(CONTRACTS_DIR / "license_prism.docx")


def create_msa():
    doc = Document()
    doc.add_heading("Professional Services MSA", 0)
    add_para(doc, "This Master Services Agreement is entered into as of February 15, 2025, by and between Horizon Digital Services (\"Provider\") and Client (\"Client\").")

    add_heading(doc, "1. Scope")
    add_para(doc, "Provider shall perform professional services as described in individually executed Statements of Work (\"SOWs\") under this MSA.")

    add_heading(doc, "2. Term")
    add_para(doc, "This MSA is effective from February 15, 2025 through February 14, 2027, with automatic annual renewal unless either party provides 90 days written notice.")

    add_heading(doc, "3. Statements of Work")
    add_para(doc, "Client shall review and approve or reject each Statement of Work within 10 business days of receipt. SOWs shall be deemed approved if Client does not respond within 10 business days.")

    add_heading(doc, "4. Change Management")
    add_para(doc, "Any changes to scope, timeline, or fees must follow the formal change order process with written approval from both parties before implementation.")

    add_heading(doc, "5. Acceptance")
    add_para(doc, "Client shall complete acceptance testing of each deliverable within 15 business days of delivery. Failure to respond within the acceptance period constitutes acceptance.")

    add_heading(doc, "6. Payment Terms")
    add_para(doc, "Client shall pay all undisputed invoices within 30 days of receipt. Late payment incurs 1% monthly interest on the outstanding amount.")

    add_heading(doc, "7. Confidentiality")
    add_para(doc, "Both parties shall maintain confidentiality of all proprietary information exchanged under this MSA in accordance with the mutual NDA executed between the parties. Breach of confidentiality subjects the breaching party to injunctive relief and damages.")

    add_heading(doc, "8. Indemnification")
    add_para(doc, "Each party shall indemnify the other against third-party claims arising from the indemnifying party's negligence or willful misconduct.")

    doc.save(CONTRACTS_DIR / "msa_horizon.docx")


def create_contractor_agreement():
    doc = Document()
    doc.add_heading("Independent Contractor Agreement", 0)
    add_para(doc, "This Agreement is entered into as of May 1, 2025, by and between Client (\"Company\") and Jordan Chen (\"Contractor\").")

    add_heading(doc, "1. Engagement")
    add_para(doc, "Company engages Contractor as an independent contractor to perform software development services as described in the attached Statement of Work.")

    add_heading(doc, "2. Term")
    add_para(doc, "This Agreement is effective from May 1, 2025 through October 31, 2025. No automatic renewal.")

    add_heading(doc, "3. Reporting")
    add_para(doc, "Contractor shall submit weekly time reports by end of day each Friday, detailing hours worked and tasks completed.")

    add_heading(doc, "4. Deliverables")
    add_para(doc, "Contractor shall deliver all work product according to the milestone schedule in the attached SOW. Missed milestones may result in payment withholding for affected deliverables.")

    add_heading(doc, "5. Termination")
    add_para(doc, "Contractor shall return all Company equipment within 5 business days of engagement termination.")

    add_heading(doc, "6. Restrictive Covenants")
    add_para(doc, "Contractor shall not solicit any Company employees or contractors for 6 months following termination of this Agreement.")

    add_heading(doc, "7. Tax")
    add_para(doc, "Contractor shall provide updated W-9 form annually and as needed for tax reporting purposes. Contractor is solely responsible for all tax obligations.")

    add_heading(doc, "8. Intellectual Property")
    add_para(doc, "All work product shall be the exclusive property of Company as work-made-for-hire.")

    doc.save(CONTRACTS_DIR / "contractor_jordan_chen.docx")


def create_lease_agreement():
    doc = Document()
    doc.add_heading("Office Space Sublease Agreement", 0)
    add_para(doc, "This Sublease Agreement is entered into as of January 1, 2025, by and between Sterling Properties Ltd. (\"Landlord\") and Client (\"Tenant\").")

    add_heading(doc, "1. Premises")
    add_para(doc, "Landlord subleases to Tenant approximately 3,500 square feet of office space located at 200 Innovation Drive, Suite 400.")

    add_heading(doc, "2. Term")
    add_para(doc, "This Sublease is effective from January 1, 2025 through December 31, 2027.")

    add_heading(doc, "3. Rent")
    add_para(doc, "Tenant shall pay monthly rent of $8,500 by the 1st of each month. Rent received after the 5th is subject to a late fee of $250 plus 5% of monthly rent for each day beyond the 5th.")

    add_heading(doc, "4. Insurance")
    add_para(doc, "Tenant shall provide proof of renter's insurance with minimum $1,000,000 liability coverage within 30 days of lease commencement and upon each annual renewal. Failure to maintain insurance is grounds for lease termination.")

    add_heading(doc, "5. Maintenance")
    add_para(doc, "Landlord shall respond to all maintenance requests within 48 hours and complete non-emergency repairs within 10 business days.")

    add_heading(doc, "6. Use of Premises")
    add_para(doc, "Tenant shall comply with all building rules, regulations, and applicable local ordinances regarding use of the premises.")

    add_heading(doc, "7. Termination")
    add_para(doc, "Either party must provide 90 days written notice of intent to terminate or not renew the lease. Failure to provide timely notice results in automatic 12-month renewal.")

    add_heading(doc, "8. Security Deposit")
    add_para(doc, "Tenant shall pay a security deposit equal to two months' rent ($17,000) upon execution of this Sublease.")

    doc.save(CONTRACTS_DIR / "lease_sterling.docx")


if __name__ == "__main__":
    create_saas_agreement()
    create_vendor_agreement()
    create_dpa()
    create_consulting_agreement()
    create_license_agreement()
    create_msa()
    create_contractor_agreement()
    create_lease_agreement()
    print(f"Generated 8 sample contracts in {CONTRACTS_DIR}")
